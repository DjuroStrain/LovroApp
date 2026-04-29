import sys
import os
import cv2
from datetime import datetime
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout,
    QHBoxLayout, QPushButton, QLabel, QMessageBox
)
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QImage, QPixmap
from docx import Document
from lxml import etree


CAPTURES_DIR = "captures"
MAX_PHOTOS = 4
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "template", "IL 30-1_r6_PT Record.docx")

# Anchor positions from r5 template (left → right), all values in EMU
_IMG_SLOTS = [
    {'posH': -8255,   'cx': 929640},
    {'posH': 1050925, 'cx': 929640},
    {'posH': 2294255, 'cx': 929640},
    {'posH': 3521075, 'cx': 929640},
]
_IMG_CY    = 1402080
_IMG_POS_V = 24765

_WP   = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
_WP14 = 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing'
_A    = 'http://schemas.openxmlformats.org/drawingml/2006/main'
_PIC  = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
_R    = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
_A14  = 'http://schemas.microsoft.com/office/drawing/2010/main'
_W    = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _build_anchor(rId, idx, slot):
    cx, cy = slot['cx'], _IMG_CY
    xml = (
        f'<wp:anchor'
        f' xmlns:wp="{_WP}" xmlns:wp14="{_WP14}"'
        f' xmlns:a="{_A}" xmlns:pic="{_PIC}"'
        f' xmlns:r="{_R}" xmlns:a14="{_A14}"'
        f' distT="0" distB="0" distL="114300" distR="114300"'
        f' simplePos="0" relativeHeight="{251659264 + idx * 1024}"'
        f' behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1"'
        f' wp14:anchorId="{0x15000000 + idx:08X}"'
        f' wp14:editId="{0x10000000 + idx:08X}">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column">'
        f'  <wp:posOffset>{slot["posH"]}</wp:posOffset>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph">'
        f'  <wp:posOffset>{_IMG_POS_V}</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{100 + idx}" name="Photo {idx + 1}"/>'
        f'<wp:cNvGraphicFramePr>'
        f'  <a:graphicFrameLocks noChangeAspect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'<a:graphic>'
        f'  <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'    <pic:pic>'
        f'      <pic:nvPicPr>'
        f'        <pic:cNvPr id="0" name="Photo {idx + 1}"/>'
        f'        <pic:cNvPicPr>'
        f'          <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>'
        f'        </pic:cNvPicPr>'
        f'      </pic:nvPicPr>'
        f'      <pic:blipFill>'
        f'        <a:blip r:embed="{rId}">'
        f'          <a:extLst>'
        f'            <a:ext uri="{{28A0092B-C50C-407E-A947-70E740481C1C}}">'
        f'              <a14:useLocalDpi val="0"/>'
        f'            </a:ext>'
        f'          </a:extLst>'
        f'        </a:blip>'
        f'        <a:stretch><a:fillRect/></a:stretch>'
        f'      </pic:blipFill>'
        f'      <pic:spPr bwMode="auto">'
        f'        <a:xfrm>'
        f'          <a:off x="0" y="0"/>'
        f'          <a:ext cx="{cx}" cy="{cy}"/>'
        f'        </a:xfrm>'
        f'        <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'        <a:noFill/>'
        f'      </pic:spPr>'
        f'    </pic:pic>'
        f'  </a:graphicData>'
        f'</a:graphic>'
        f'<wp14:sizeRelH relativeFrom="page">'
        f'  <wp14:pctWidth>0</wp14:pctWidth>'
        f'</wp14:sizeRelH>'
        f'<wp14:sizeRelV relativeFrom="page">'
        f'  <wp14:pctHeight>0</wp14:pctHeight>'
        f'</wp14:sizeRelV>'
        f'</wp:anchor>'
    )
    return etree.fromstring(xml)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Photo Capture")
        self.setMinimumSize(800, 620)

        self.captured_paths = []
        self.camera = None
        self.timer = QTimer()

        self._init_ui()
        self._init_camera()

    def _init_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setSpacing(12)
        layout.setContentsMargins(16, 16, 16, 16)

        self.preview_label = QLabel("Initializing camera...")
        self.preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.preview_label.setMinimumHeight(480)
        self.preview_label.setStyleSheet("background: #111; color: #aaa; font-size: 14px;")
        layout.addWidget(self.preview_label)

        controls = QHBoxLayout()

        self.counter_label = QLabel(f"Photos: 0 / {MAX_PHOTOS}")
        self.counter_label.setStyleSheet("font-size: 16px; font-weight: bold;")
        controls.addWidget(self.counter_label)

        controls.addStretch()

        self.capture_btn = QPushButton("Capture Photo")
        self.capture_btn.setFixedHeight(40)
        self.capture_btn.setStyleSheet("font-size: 14px; padding: 0 20px;")
        self.capture_btn.clicked.connect(self.capture_photo)
        controls.addWidget(self.capture_btn)

        self.save_btn = QPushButton("Save Document")
        self.save_btn.setFixedHeight(40)
        self.save_btn.setStyleSheet("font-size: 14px; padding: 0 20px;")
        self.save_btn.setEnabled(False)
        self.save_btn.clicked.connect(self.save_document)
        controls.addWidget(self.save_btn)

        layout.addLayout(controls)

    def _init_camera(self):
        self.camera = cv2.VideoCapture(0)
        if not self.camera.isOpened():
            QMessageBox.critical(self, "Camera Error", "No webcam detected. Please connect a USB camera and restart the app.")
            self.capture_btn.setEnabled(False)
            return

        self.timer.timeout.connect(self._update_frame)
        self.timer.start(33)

    def _update_frame(self):
        if self.camera is None or not self.camera.isOpened():
            return
        ret, frame = self.camera.read()
        if not ret:
            return
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        h, w, ch = frame_rgb.shape
        img = QImage(frame_rgb.data, w, h, ch * w, QImage.Format.Format_RGB888)
        pixmap = QPixmap.fromImage(img).scaled(
            self.preview_label.width(),
            self.preview_label.height(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation,
        )
        self.preview_label.setPixmap(pixmap)

    def capture_photo(self):
        if self.camera is None or not self.camera.isOpened():
            QMessageBox.warning(self, "Camera Error", "Camera is not available.")
            return

        ret, frame = self.camera.read()
        if not ret:
            QMessageBox.warning(self, "Capture Failed", "Failed to capture image from camera.")
            return

        os.makedirs(CAPTURES_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        filename = os.path.join(CAPTURES_DIR, f"photo_{timestamp}.jpg")
        cv2.imwrite(filename, frame)
        self.captured_paths.append(os.path.abspath(filename))

        count = len(self.captured_paths)
        self.counter_label.setText(f"Photos: {count} / {MAX_PHOTOS}")

        if count >= MAX_PHOTOS:
            self.capture_btn.setEnabled(False)
            self.save_btn.setEnabled(True)

    def save_document(self):
        doc = Document(TEMPLATE_PATH)

        rIds = []
        for path in self.captured_paths:
            rId, _ = doc.part.get_or_add_image(path)
            rIds.append(rId)

        # Find the empty paragraph after "Slika/Image:" in the table
        target_para = None
        for table in doc.tables:
            if target_para:
                break
            for row in table.rows:
                if target_para:
                    break
                for cell in row.cells:
                    paras = cell.paragraphs
                    for i, para in enumerate(paras):
                        if 'Slika' in para.text and i + 1 < len(paras):
                            target_para = paras[i + 1]
                            break
                    if target_para:
                        break

        if target_para is None:
            QMessageBox.critical(self, "Template Error", "Could not find image placeholder in template.")
            return

        for idx, (rId, slot) in enumerate(zip(rIds, _IMG_SLOTS)):
            anchor_el = _build_anchor(rId, idx, slot)
            r_el = etree.SubElement(target_para._element, f'{{{_W}}}r')
            drawing_el = etree.SubElement(r_el, f'{{{_W}}}drawing')
            drawing_el.append(anchor_el)

        docs_folder = os.path.expanduser("~/Documents")
        os.makedirs(docs_folder, exist_ok=True)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        out_path = os.path.join(docs_folder, f"PT_Record_{timestamp}.docx")

        try:
            doc.save(out_path)
        except Exception as e:
            QMessageBox.critical(self, "Save Error", f"Could not save document:\n{e}")
            return

        QMessageBox.information(self, "Saved", f"Document saved to:\n{out_path}")

        self.captured_paths = []
        self.counter_label.setText(f"Photos: 0 / {MAX_PHOTOS}")
        self.capture_btn.setEnabled(True)
        self.save_btn.setEnabled(False)

    def closeEvent(self, event):
        self.timer.stop()
        if self.camera and self.camera.isOpened():
            self.camera.release()
        event.accept()


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
